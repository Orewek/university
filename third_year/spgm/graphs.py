# type: ignore
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import re
import matplotlib.backends.backend_pdf
from docx import Document
from docx.shared import Inches
import io


# Настройка стиля графиков
plt.style.use('default')
plt.rcParams['figure.figsize'] = [12, 8]
plt.rcParams['font.size'] = 12
plt.rcParams['font.family'] = 'DejaVu Sans'


def parse_md_file(filename: str):
    """
    Парсинг файла с результатами молекулярной динамики
    """
    steps = []
    data = {
        'Ekin': [], 'Eterm': [], 'Epot': [], 'Eint': [],
        'E': [], 'T': [], 'P': [],
        'r1': [], 'r2': [], 'r12_abs': [],
        'U12': [], 'F12': [], 'F1': [],
    }

    with open(filename, 'r', encoding='utf-8') as file:
        content = file.read()

    # Разделяем по шагам
    step_blocks = content.split('Step =')

    for block in step_blocks[1:]:  # Пропускаем первый пустой элемент
        lines = block.strip().split('\n')
        # Извлекаем номер шага
        step_match = re.match(r'(\d+)', lines[0])
        if step_match:
            steps.append(int(step_match.group(1)))

            # Инициализируем значения для шага
            [data[key].append(np.nan) for key in data]

        # Парсим параметры
        for line in lines[1:]:
            line = line.strip()
            if '=' not in line:
                continue

            parts = line.split('=')
            if len(parts) != 2:
                continue

            key = parts[0].strip()
            value = parts[1].strip()

            if key not in data:
                continue

            try:
                data[key][-1] = float(value)

            except ValueError:
                pass
    return steps, data


def create_energy_plot(steps, data):
    """
    Создает график зависимостей энергий от номера шага
    """
    _, ax = plt.subplots(figsize=(10, 6))

    # Построение графиков энергий
    ax.plot(
        steps, data['Eterm'],
        label='Eterm (Тепловая энергия)',
        linewidth=2,
        color='blue',
        alpha=0.8,
    )
    ax.plot(
        steps,
        data['Epot'],
        label='Epot (Потенциальная энергия)',
        linewidth=2,
        color='red',
        alpha=0.8,
    )
    ax.plot(
        steps,
        data['E'],
        label='E (Полная энергия)',
        linewidth=2,
        color='green',
        alpha=0.8,
    )

    ax.set_xlabel('Номер шага', fontsize=14)
    ax.set_ylabel('Энергия', fontsize=14)
    ax.set_title('Зависимость энергий от номера шага', fontsize=16)
    ax.legend(fontsize=12)
    ax.grid(True, alpha=0.3)
    ax.set_xlim(0, max(steps))

    plt.tight_layout()
    plt.savefig('energy_plot.png', dpi=300, bbox_inches='tight')
    plt.show()


def create_temperature_plot(steps, data):
    """
    Создает график зависимости температуры от номера шага
    """
    _, ax = plt.subplots(figsize=(10, 6))

    ax.plot(
        steps,
        data['T'],
        label='T (Температура)',
        linewidth=2,
        color='orange',
        alpha=0.8,
    )

    ax.set_xlabel('Номер шага', fontsize=14)
    ax.set_ylabel('Температура', fontsize=14)
    ax.set_title('Зависимость температуры от номера шага', fontsize=16)
    ax.legend(fontsize=12)
    ax.grid(True, alpha=0.3)
    ax.set_xlim(600, 1000)

    plt.tight_layout()
    plt.savefig('temperature_plot.png', dpi=300, bbox_inches='tight')
    plt.show()


def create_pressure_plot(steps, data):
    """
    Создает график зависимости давления от номера шага
    """
    _, ax = plt.subplots(figsize=(10, 6))

    ax.plot(
        steps,
        data['P'],
        label='P (Давление)',
        linewidth=2,
        color='purple',
        alpha=0.8,
    )

    ax.set_xlabel('Номер шага', fontsize=14)
    ax.set_ylabel('Давление', fontsize=14)
    ax.set_title('Зависимость давления от номера шага', fontsize=16)
    ax.legend(fontsize=12)
    ax.grid(True, alpha=0.3)
    ax.set_xlim(600, 1000)

    plt.tight_layout()
    plt.savefig('pressure_plot.png', dpi=300, bbox_inches='tight')
    plt.show()


def create_comprehensive_plot(steps, data):
    """
    Создает комплексный график со всеми параметрами
    """
    _, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(15, 10))

    # 1. Энергии
    ax1.plot(steps, data['Eterm'], label='Eterm', linewidth=2)
    ax1.plot(steps, data['Epot'], label='Epot', linewidth=2)
    ax1.plot(steps, data['E'], label='E', linewidth=2)
    ax1.set_xlabel('Номер шага')
    ax1.set_ylabel('Энергия')
    ax1.set_title('Энергии системы')
    ax1.legend()
    ax1.grid(True, alpha=0.3)

    # 2. Температура
    ax2.plot(steps, data['T'], label='T', color='red', linewidth=2)
    ax2.set_xlabel('Номер шага')
    ax2.set_ylabel('Температура')
    ax2.set_title('Температура системы')
    ax2.legend()
    ax2.grid(True, alpha=0.3)

    # 3. Давление
    ax3.plot(steps, data['P'], label='P', color='green', linewidth=2)
    ax3.set_xlabel('Номер шага')
    ax3.set_ylabel('Давление')
    ax3.set_title('Давление системы')
    ax3.legend()
    ax3.grid(True, alpha=0.3)

    # 4. Все параметры (нормализованные)
    ax4.plot(
        steps,
        (data['E'] - np.mean(data['E'])) / np.std(data['E']),
        label='E (норм.)',
        alpha=0.7,
    )
    ax4.plot(
        steps,
        (data['T'] - np.mean(data['T'])) / np.std(data['T']),
        label='T (норм.)',
        alpha=0.7,
    )
    ax4.plot(
        steps,
        (data['P'] - np.mean(data['P'])) / np.std(data['P']),
        label='P (норм.)',
        alpha=0.7,
    )
    ax4.set_xlabel('Номер шага')
    ax4.set_ylabel('Нормализованные значения')
    ax4.set_title('Нормализованные параметры')
    ax4.legend()
    ax4.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig('comprehensive_plot.png', dpi=300, bbox_inches='tight')
    plt.show()


def main():
    filename: str = "Charykov.txt"

    print("Чтение данных из файла...")
    steps, data = parse_md_file(filename)
    print(f"Успешно прочитано {len(steps)} шагов")
    # Создаем отдельные графики как в вашем документе

    print("Создание PDF файла...")
    create_plots_pdf(steps, data)
    create_plots_docx(steps, data)

    # create_energy_plot(steps, data)
    # create_temperature_plot(steps, data)
    # create_pressure_plot(steps, data)
    # create_comprehensive_plot(steps, data)
    # Сохраняем данные в CSV
    df = pd.DataFrame(data)
    df['Step'] = steps
    df.to_csv('md_results_data.csv', index=False)


def create_plots_pdf(steps, data):
    """
    Создает PDF файл со всеми графиками
    """
    with matplotlib.backends.backend_pdf.PdfPages('Charykov_MD_Results.pdf') as pdf:

        # 1. График энергий
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(
            steps, data['Eterm'],
            label='Eterm (Тепловая энергия)',
            linewidth=2,
            color='blue',
            alpha=0.8,
        )
        ax.plot(
            steps,
            data['Epot'],
            label='Epot (Потенциальная энергия)',
            linewidth=2,
            color='green',
            alpha=0.8,
        )
        ax.plot(
            steps,
            data['E'],
            label='E (Полная энергия)',
            linewidth=2,
            color='red',
            alpha=0.8,
        )
        ax.set_xlabel('Номер шага', fontsize=14)
        ax.set_ylabel('Энергия', fontsize=14)
        ax.set_title('Зависимость энергий от номера шага', fontsize=16)
        ax.legend(fontsize=12)
        ax.grid(True, alpha=0.3)
        ax.set_xlim(0, max(steps))
        plt.tight_layout()
        pdf.savefig(fig)
        plt.close(fig)

        # 2. График температуры
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(steps, data['T'], label='T', color='red', linewidth=2)
        ax.set_xlabel('Номер шага')
        ax.set_ylabel('Температура')
        ax.set_title('Температура системы')
        ax.legend(fontsize=12)
        ax.grid(True, alpha=0.3)

        plt.tight_layout()
        pdf.savefig(fig)
        plt.close(fig)

        # 3. График давления
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(steps, data['P'], label='P', color='blue', linewidth=2)
        ax.set_xlabel('Номер шага')
        ax.set_ylabel('Давление')
        ax.set_title('Давление системы')
        ax.legend(fontsize=12)
        ax.grid(True, alpha=0.3)

        plt.tight_layout()
        pdf.savefig(fig)
        plt.close(fig)


def create_plots_docx(steps, data):
    """
    Создает DOCX файл со всеми графиками
    """
    # Создаем новый документ
    doc = Document()

    # Добавляем заголовок
    doc.add_heading('Результаты молекулярно-динамического моделирования', 0)

    # 1. График энергий
    doc.add_heading('График энергий', level=1)

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(
        steps, data['Eterm'],
        label='Eterm (Тепловая энергия)',
        linewidth=2,
        color='blue',
        alpha=0.8,
    )
    ax.plot(
        steps,
        data['Epot'],
        label='Epot (Потенциальная энергия)',
        linewidth=2,
        color='green',
        alpha=0.8,
    )
    ax.plot(
        steps,
        data['E'],
        label='E (Полная энергия)',
        linewidth=2,
        color='red',
        alpha=0.8,
    )
    ax.set_xlabel('Номер шага', fontsize=14)
    ax.set_ylabel('Энергия', fontsize=14)
    ax.set_title('Зависимость энергий от номера шага', fontsize=16)
    ax.legend(fontsize=12)
    ax.grid(True, alpha=0.3)
    ax.set_xlim(0, max(steps))
    plt.tight_layout()

    # Сохраняем график в буфер и добавляем в документ
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    plt.close(fig)

    # Добавляем описание
    doc.add_paragraph('График показывает изменение различных видов энергии системы во времени')

    # 2. График температуры
    doc.add_heading('График температуры', level=1)

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(steps, data['T'], label='T', color='red', linewidth=2)
    ax.set_xlabel('Номер шага')
    ax.set_ylabel('Температура')
    ax.set_title('Температура системы')
    ax.legend(fontsize=12)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()

    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    plt.close(fig)

    doc.add_paragraph('График показывает изменение температуры системы во времени.')

    # 3. График давления
    doc.add_heading('График давления', level=1)

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(steps, data['P'], label='P', color='blue', linewidth=2)
    ax.set_xlabel('Номер шага')
    ax.set_ylabel('Давление')
    ax.set_title('Давление системы')
    ax.legend(fontsize=12)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()

    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    plt.close(fig)

    doc.add_paragraph('График показывает изменение давления системы во времени.')

    # Сохраняем документ
    doc.save('Charykov_MD_Results.docx')

    print("DOCX файл с графиками успешно создан: Charykov_MD_Results.docx")


if __name__ == "__main__":
    main()
